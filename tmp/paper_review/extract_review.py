from __future__ import annotations

from collections import Counter
from pathlib import Path
import json
import re

from docx import Document
from pypdf import PdfReader


ROOT = Path.cwd()
OUT = ROOT / "tmp" / "paper_review"
DOCX_PATH = ROOT / "余涛初稿改3.docx"
PDF_PATH = ROOT / "余涛初稿.pdf"


def extract_docx() -> dict:
    doc = Document(DOCX_PATH)
    paras: list[dict] = []
    for idx, paragraph in enumerate(doc.paragraphs, 1):
        text = paragraph.text.replace("\u00a0", " ").strip()
        style = paragraph.style.name if paragraph.style else ""
        if text:
            paras.append({"idx": idx, "style": style, "text": text})

    with (OUT / "docx_paragraphs.txt").open("w", encoding="utf-8") as handle:
        for item in paras:
            handle.write(f"[{item['idx']:04d}] ({item['style']}) {item['text']}\n")

    with (OUT / "docx_tables.txt").open("w", encoding="utf-8") as handle:
        for table_idx, table in enumerate(doc.tables, 1):
            cols = len(table.columns) if table.rows else 0
            handle.write(f"# Table {table_idx}, rows={len(table.rows)}, cols={cols}\n")
            for row_idx, row in enumerate(table.rows, 1):
                cells = [" ".join(cell.text.split()) for cell in row.cells]
                handle.write(f"{row_idx}: " + " | ".join(cells) + "\n")
            handle.write("\n")

    heading_like: list[dict] = []
    heading_re = re.compile(
        r"^(第[一二三四五六七八九十]+章|[1-9](?:\.[0-9]+)*\s|摘\s*要|Abstract|"
        r"致\s*谢|参考文献|附录|结论)"
    )
    for item in paras:
        if heading_re.match(item["text"]):
            heading_like.append(item)

    with (OUT / "docx_headings.txt").open("w", encoding="utf-8") as handle:
        for item in heading_like:
            handle.write(f"[{item['idx']:04d}] ({item['style']}) {item['text']}\n")

    full_text = "\n".join(item["text"] for item in paras)
    patterns = {
        "TODO/占位": r"TODO|待补|XXX|xx|TBD|补充|此处|\[\?\]|？？",
        "连续英文破折/占位线": r"-{4,}|_{4,}",
        "空引用": r"\[\s*\]",
        "英文中文粘连": r"[A-Za-z]{2,}[\u4e00-\u9fff]|[\u4e00-\u9fff][A-Za-z]{2,}",
        "全角括号内纯英文": r"（[A-Za-z0-9 ,;:/_.-]+）",
    }
    pattern_counts = {name: len(re.findall(regex, full_text)) for name, regex in patterns.items()}
    styles = Counter(item["style"] for item in paras)

    summary = {
        "docx": str(DOCX_PATH),
        "paragraph_count_nonempty": len(paras),
        "table_count": len(doc.tables),
        "section_count": len(doc.sections),
        "characters": len(full_text),
        "style_top": styles.most_common(20),
        "pattern_counts": pattern_counts,
    }
    (OUT / "docx_summary.json").write_text(
        json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return summary


def extract_pdf() -> dict:
    reader = PdfReader(PDF_PATH)
    page_texts = []
    for page_idx, page in enumerate(reader.pages, 1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:  # pragma: no cover - best-effort extraction
            text = f"<EXTRACT_ERROR {exc}>"
        page_texts.append((page_idx, text))

    with (OUT / "pdf_text_by_page.txt").open("w", encoding="utf-8", errors="replace") as handle:
        for page_idx, text in page_texts:
            handle.write(f"\n\n===== PAGE {page_idx} =====\n{text}\n")

    summary = {
        "pdf": str(PDF_PATH),
        "page_count": len(reader.pages),
        "metadata": {str(k): str(v) for k, v in (reader.metadata or {}).items()},
        "extractable_chars": sum(len(text) for _, text in page_texts),
    }
    (OUT / "pdf_summary.json").write_text(
        json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return summary


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    result = {"docx": extract_docx(), "pdf": extract_pdf()}
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
