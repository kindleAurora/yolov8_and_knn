from __future__ import annotations

import argparse
import difflib
import hashlib
import json
import re
import zipfile
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from PIL import Image, ImageDraw, ImageFont


HEADING_RE = re.compile(
    r"^("
    r"\u6458\s*\u8981|Abstract|\u76ee\s*\u5f55|"
    r"\u7b2c[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+\u7ae0|"
    r"[1-9](?:\.[0-9]+)*\s+|"
    r"\u53c2\u8003\u6587\u732e|\u81f4\s*\u8c22|\u9644\u5f55"
    r")"
)

PLACEHOLDER_RE = re.compile(
    r"TODO|TBD|XXX|\u5f85\u8865|\u6b64\u5904|\u8865\u5145|\[\?\]|\?\?|\uff1f\uff1f",
    re.IGNORECASE,
)
CITATION_RE = re.compile(r"\[(\d+(?:\s*[-,，]\s*\d+)*)\]")
FIG_RE = re.compile(r"\u56fe\s*([0-9]+)\s*[-\uff0d\u2011\u2013\u2014]\s*([0-9]+)")
TABLE_RE = re.compile(r"\u8868\s*([0-9]+)\s*[-\uff0d\u2011\u2013\u2014]\s*([0-9]+)")


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def para_records(doc: Document) -> list[dict]:
    records = []
    for idx, para in enumerate(doc.paragraphs, 1):
        text = para.text.replace("\u00a0", " ").strip()
        if text:
            records.append(
                {
                    "idx": idx,
                    "style": para.style.name if para.style else "",
                    "text": text,
                }
            )
    return records


def table_records(doc: Document) -> list[dict]:
    out = []
    for table_idx, table in enumerate(doc.tables, 1):
        rows = []
        for row in table.rows:
            rows.append([" ".join(cell.text.split()) for cell in row.cells])
        out.append(
            {
                "idx": table_idx,
                "rows": len(table.rows),
                "cols": len(table.columns) if table.rows else 0,
                "preview": rows[:4],
            }
        )
    return out


def zip_counts(path: Path) -> dict:
    with zipfile.ZipFile(path) as zf:
        names = zf.namelist()
        xml_text = "\n".join(
            zf.read(name).decode("utf-8", errors="ignore")
            for name in names
            if name.endswith(".xml") and name.startswith("word/")
        )
    return {
        "media_files": [name for name in names if name.startswith("word/media/")],
        "has_comments": "word/comments.xml" in names,
        "tracked_insertions": xml_text.count("<w:ins"),
        "tracked_deletions": xml_text.count("<w:del"),
        "field_codes": xml_text.count("<w:fldChar"),
    }


def image_inventory(path: Path) -> list[dict]:
    items = []
    with zipfile.ZipFile(path) as zf:
        for name in zf.namelist():
            if not name.startswith("word/media/"):
                continue
            try:
                with zf.open(name) as f:
                    image = Image.open(f)
                    items.append(
                        {
                            "name": name,
                            "width": image.width,
                            "height": image.height,
                            "mode": image.mode,
                        }
                    )
            except Exception as exc:
                items.append({"name": name, "error": str(exc)})
    return items


def normalize_caption(kind: str, match: re.Match) -> str:
    return f"{kind}{match.group(1)}-{match.group(2)}"


def find_contexts(records: list[dict], regex: re.Pattern, limit: int = 30) -> list[dict]:
    contexts = []
    for item in records:
        if regex.search(item["text"]):
            contexts.append(item)
            if len(contexts) >= limit:
                break
    return contexts


def reference_entries(records: list[dict]) -> list[dict]:
    start = None
    for pos, item in enumerate(records):
        if re.match(r"^\u53c2\u8003\u6587\u732e$", item["text"]):
            start = pos + 1
            break
    if start is None:
        return []
    entries = []
    for item in records[start:]:
        if re.match(r"^(\u81f4\s*\u8c22|\u9644\u5f55)", item["text"]):
            break
        if re.match(r"^(\[\d+\]|\d+[\.\u3001])", item["text"]):
            entries.append(item)
    return entries


def citation_first_order(text: str) -> list[int]:
    seen = []
    for match in CITATION_RE.finditer(text):
        first = int(re.split(r"[-,，]", match.group(1))[0].strip())
        if first not in seen:
            seen.append(first)
    return seen


def expected_gaps(nums: list[str]) -> dict[str, list[int]]:
    by_chapter: dict[str, set[int]] = defaultdict(set)
    for item in nums:
        m = re.match(r".*?(\d+)-(\d+)", item)
        if m:
            by_chapter[m.group(1)].add(int(m.group(2)))
    gaps = {}
    for chapter, values in by_chapter.items():
        if not values:
            continue
        missing = [n for n in range(1, max(values) + 1) if n not in values]
        if missing:
            gaps[chapter] = missing
    return gaps


def summarize_docx(path: Path) -> dict:
    doc = Document(path)
    records = para_records(doc)
    tables = table_records(doc)
    full_text = "\n".join(item["text"] for item in records)
    table_text = "\n".join(" ".join(" ".join(row) for row in table["preview"]) for table in tables)
    all_text = full_text + "\n" + table_text

    headings = [item for item in records if HEADING_RE.match(item["text"])]
    fig_refs = [normalize_caption("\u56fe", m) for m in FIG_RE.finditer(all_text)]
    table_refs = [normalize_caption("\u8868", m) for m in TABLE_RE.finditer(all_text)]
    fig_caption_items = [item for item in records if re.match(r"^\u56fe\s*\d+", item["text"])]
    table_caption_items = [item for item in records if re.match(r"^\u8868\s*\d+", item["text"])]
    citations = citation_first_order(full_text)
    refs = reference_entries(records)
    zinfo = zip_counts(path)
    images = image_inventory(path)

    sections = []
    for idx, section in enumerate(doc.sections, 1):
        sections.append(
            {
                "idx": idx,
                "page_width_cm": round(section.page_width.cm, 2),
                "page_height_cm": round(section.page_height.cm, 2),
                "top_margin_cm": round(section.top_margin.cm, 2),
                "bottom_margin_cm": round(section.bottom_margin.cm, 2),
                "left_margin_cm": round(section.left_margin.cm, 2),
                "right_margin_cm": round(section.right_margin.cm, 2),
            }
        )

    return {
        "path": str(path),
        "size": path.stat().st_size,
        "sha256": sha256(path),
        "paragraph_count_nonempty": len(records),
        "character_count": len(full_text),
        "table_count": len(tables),
        "section_count": len(doc.sections),
        "style_top": Counter(item["style"] for item in records).most_common(20),
        "headings": headings,
        "tables": tables,
        "sections": sections,
        "figures": {
            "caption_items": fig_caption_items,
            "ref_counter": Counter(fig_refs).most_common(),
            "duplicates": [k for k, v in Counter(fig_refs).items() if v > 1],
            "gaps": expected_gaps(sorted(set(fig_refs))),
        },
        "tables_caption": {
            "caption_items": table_caption_items,
            "ref_counter": Counter(table_refs).most_common(),
            "duplicates": [k for k, v in Counter(table_refs).items() if v > 1],
            "gaps": expected_gaps(sorted(set(table_refs))),
        },
        "references": {
            "entry_count": len(refs),
            "entries": refs,
            "citation_first_order": citations,
            "citation_max": max(citations) if citations else None,
            "citation_count": len(re.findall(CITATION_RE, full_text)),
        },
        "placeholders": find_contexts(records, PLACEHOLDER_RE),
        "term_counts": {
            "YOLOv8": len(re.findall(r"YOLOv8", all_text, re.I)),
            "YOLO v8": len(re.findall(r"YOLO\s+v8", all_text, re.I)),
            "DeepSORT": len(re.findall(r"DeepSORT", all_text, re.I)),
            "Deep SORT": len(re.findall(r"Deep\s+SORT", all_text, re.I)),
            "KNN": len(re.findall(r"\bKNN\b", all_text, re.I)),
            "K-NN": len(re.findall(r"K[-\s]NN", all_text, re.I)),
        },
        "zip": {
            "media_count": len(zinfo["media_files"]),
            "has_comments": zinfo["has_comments"],
            "tracked_insertions": zinfo["tracked_insertions"],
            "tracked_deletions": zinfo["tracked_deletions"],
            "field_codes": zinfo["field_codes"],
        },
        "images": images,
        "first_paragraphs": records[:40],
        "last_paragraphs": records[-30:],
    }


def compare_text(primary: dict, other_path: Path) -> dict:
    other = summarize_docx(other_path)
    a = [item["text"] for item in primary["headings"]]
    b = [item["text"] for item in other["headings"]]
    heading_diff = list(difflib.unified_diff(b, a, fromfile=str(other_path), tofile=primary["path"], lineterm=""))
    return {
        "other_path": str(other_path),
        "other_sha256": other["sha256"],
        "same_binary": primary["sha256"] == other["sha256"],
        "paragraph_count_other": other["paragraph_count_nonempty"],
        "character_count_other": other["character_count"],
        "heading_diff": heading_diff[:120],
    }


def natural_page_key(path: Path) -> int:
    m = re.search(r"page-(\d+)\.png$", path.name)
    return int(m.group(1)) if m else 0


def make_contact_sheets(render_dir: Path, out_dir: Path, per_sheet: int = 12) -> list[str]:
    pages = sorted(render_dir.glob("page-*.png"), key=natural_page_key)
    if not pages:
        return []
    out_dir.mkdir(parents=True, exist_ok=True)
    sheets = []
    thumb_w = 240
    label_h = 28
    cols = 4
    rows = (per_sheet + cols - 1) // cols
    font = ImageFont.load_default()
    for sheet_idx in range(0, len(pages), per_sheet):
        chunk = pages[sheet_idx : sheet_idx + per_sheet]
        thumbs = []
        for page in chunk:
            im = Image.open(page).convert("RGB")
            thumb_h = round(im.height * thumb_w / im.width)
            im = im.resize((thumb_w, thumb_h))
            canvas = Image.new("RGB", (thumb_w, thumb_h + label_h), "white")
            canvas.paste(im, (0, label_h))
            draw = ImageDraw.Draw(canvas)
            draw.text((6, 8), page.stem, fill="black", font=font)
            thumbs.append(canvas)
        cell_h = max(t.height for t in thumbs)
        sheet = Image.new("RGB", (cols * thumb_w, rows * cell_h), "white")
        for i, thumb in enumerate(thumbs):
            x = (i % cols) * thumb_w
            y = (i // cols) * cell_h
            sheet.paste(thumb, (x, y))
        out_path = out_dir / f"contact-{sheet_idx // per_sheet + 1}.png"
        sheet.save(out_path)
        sheets.append(str(out_path))
    return sheets


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", required=True)
    parser.add_argument("--compare")
    parser.add_argument("--render-dir")
    parser.add_argument("--out-dir", default="tmp/paper_review/latest")
    args = parser.parse_args()

    out_dir = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    primary = summarize_docx(Path(args.docx))
    result = {"primary": primary}
    if args.compare:
        result["compare"] = compare_text(primary, Path(args.compare))
    if args.render_dir:
        result["contact_sheets"] = make_contact_sheets(Path(args.render_dir), out_dir / "contact_sheets")

    (out_dir / "review_summary.json").write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    (out_dir / "headings.txt").write_text(
        "\n".join(f"[{h['idx']:04d}] ({h['style']}) {h['text']}" for h in primary["headings"]),
        encoding="utf-8",
    )
    (out_dir / "placeholders.txt").write_text(
        "\n".join(f"[{p['idx']:04d}] ({p['style']}) {p['text']}" for p in primary["placeholders"]),
        encoding="utf-8",
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
