from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "output" / "doc"
DOCX_PATH = OUT_DIR / "融合KNN与YOLOv8的智能牧场牲畜计数与行为监测系统_论文初稿.docx"
MD_PATH = OUT_DIR / "融合KNN与YOLOv8的智能牧场牲畜计数与行为监测系统_论文初稿.md"


def set_cell_text(cell, text: str) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(10.5)


def set_style_font(style, name: str, east_asia: str, size: int, bold: bool = False) -> None:
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    style.font.size = Pt(size)
    style.font.bold = bold


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_paragraph(doc: Document, text: str = "", style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.first_line_indent = Cm(0.74) if style is None else None
    paragraph.paragraph_format.line_spacing = 1.5
    if text:
        run = paragraph.add_run(text)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)
    return paragraph


def add_heading(doc: Document, text: str, level: int):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.line_spacing = 1.5
    for run in paragraph.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        run.font.bold = True
        run.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)
    return paragraph


def add_picture_placeholder(doc: Document, caption: str, note: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f"【{caption} 此处插入{note}】")
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    run.font.bold = True
    return paragraph


def md_heading(text: str, level: int) -> str:
    return f"{'#' * level} {text}\n\n"


def build_document() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    add_page_number(section.footer.paragraphs[0])

    set_style_font(doc.styles["Normal"], "Times New Roman", "宋体", 10)
    set_style_font(doc.styles["Heading 1"], "Times New Roman", "黑体", 16, True)
    set_style_font(doc.styles["Heading 2"], "Times New Roman", "黑体", 14, True)
    set_style_font(doc.styles["Heading 3"], "Times New Roman", "黑体", 12, True)

    md: list[str] = []

    title = "融合KNN与YOLOv8的智能牧场牲畜计数与行为监测系统"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(18)
    r.font.bold = True
    md.append(f"# {title}\n\n")

    add_heading(doc, "摘 要", 1)
    md.append(md_heading("摘 要", 1))
    abstract = (
        "随着规模化养殖和智慧农业的发展，传统依靠人工巡检的牧场管理方式逐渐难以满足实时化、精细化和低成本管理需求。"
        "针对牧场场景中牲畜数量统计不便、行为状态难以及时识别、异常情况发现滞后等问题，本文设计并实现了一种融合KNN与YOLOv8的智能牧场牲畜计数与行为监测系统。"
        "系统首先利用YOLOv8模型完成牛只目标检测与数量统计，并通过迁移学习方式提升模型在奶牛数据集上的检测效果；随后对检测框中的单牛区域进行裁剪，提取HOG姿态特征，"
        "利用KNN分类器识别卧躺、站立和行走等典型行为。针对视频检测中单帧分类易受遮挡、模糊和姿态相似影响的问题，本文进一步引入连续帧目标中心位移构成的运动信息，"
        "通过动量修正规则辅助区分站立与行走状态，并结合牧场区域标注信息对采食、饮水和休息等语义行为进行修正。实验结果表明，迁移学习后的YOLOv8检测模型在验证集上mAP50达到0.9018，"
        "mAP50-95达到0.7094；KNN行为分类模型在测试集上的整体准确率达到87.96%，其中卧躺、站立和行走三类准确率分别为92.17%、83.22%和91.67%。"
        "在此基础上，本文搭建了包含用户登录、设备管理、实时监控、行为事件、告警规则和历史分析等功能的Web监控平台，实现了从视频输入、算法推理到结果展示和告警管理的基本闭环。"
    )
    add_paragraph(doc, abstract)
    add_paragraph(doc, "关键词：智能牧场；YOLOv8；KNN；牲畜计数；行为识别；视频检测")
    md.append(abstract + "\n\n关键词：智能牧场；YOLOv8；KNN；牲畜计数；行为识别；视频检测\n\n")

    add_heading(doc, "Abstract", 1)
    md.append(md_heading("Abstract", 1))
    abstract_en = (
        "With the development of large-scale livestock farming and smart agriculture, manual inspection is no longer sufficient for real-time and refined pasture management. "
        "To address the problems of cattle counting, behavior recognition and delayed abnormal-event discovery, this thesis designs an intelligent pasture monitoring system based on YOLOv8 and KNN. "
        "YOLOv8 is first used for cattle detection and counting, and transfer learning is adopted to improve detection performance on the cattle dataset. "
        "Then each detected cattle region is cropped, HOG features are extracted, and a KNN classifier is used to recognize typical behaviors including lying, standing and walking. "
        "For video scenarios, a motion-based correction strategy is introduced by measuring the displacement of tracked object centers across consecutive frames, which reduces misclassification between standing and walking. "
        "Region-based rules are also used to refine semantic behaviors such as feeding, drinking and resting. Experimental results show that the transfer-learning YOLOv8 model achieves 0.9018 mAP50 and 0.7094 mAP50-95, "
        "while the KNN behavior classifier reaches 87.96% accuracy on the test set. Based on the algorithm pipeline, a Web monitoring platform is implemented with login, device management, real-time monitoring, behavior events, alarms and historical analysis."
    )
    add_paragraph(doc, abstract_en)
    add_paragraph(doc, "Key words: smart pasture; YOLOv8; KNN; cattle counting; behavior recognition; video detection")
    md.append(abstract_en + "\n\nKey words: smart pasture; YOLOv8; KNN; cattle counting; behavior recognition; video detection\n\n")

    # Chapter 1
    add_heading(doc, "第 1 章 绪论", 1)
    md.append(md_heading("第 1 章 绪论", 1))
    add_heading(doc, "1.1 研究背景与意义", 2)
    md.append(md_heading("1.1 研究背景与意义", 2))
    paragraphs = [
        "畜牧业是农业生产体系的重要组成部分，牛只养殖规模的扩大使牧场管理逐渐从经验式管理转向数据化、自动化和智能化管理。传统牧场通常依赖人工巡检完成牲畜计数、健康观察和异常行为判断，但人工巡检受到时间、人员经验和观察范围限制，难以实现全天候连续监测。在大规模牧场中，牲畜数量多、活动范围广、遮挡频繁，人工方式不仅成本较高，而且容易出现漏检、误判和记录不完整等问题。",
        "牛只行为与健康状态、采食饮水情况、发情状态和疾病风险密切相关。卧躺时间、行走频率、长时间停留、采食和饮水行为的变化，往往能够反映个体健康和群体管理状态。若能够通过视觉监控自动识别牲畜数量和行为类别，牧场管理人员便可以及时掌握牛群状态，并在异常发生时快速处理，从而降低人工巡检压力，提高养殖管理效率。",
        "近年来，深度学习和计算机视觉技术快速发展，目标检测算法在工业检测、交通监控、农业生产和畜牧养殖等领域得到广泛应用。YOLO系列算法具有检测速度快、部署方便和实时性较好的特点，适合用于牧场视频监控场景。与此同时，KNN等传统机器学习方法具有实现简单、训练成本低和可解释性较强的优点，可在样本规模有限的情况下完成基础行为分类。将YOLOv8目标检测能力与KNN行为分类能力结合，可以形成从目标定位、数量统计到行为识别的完整流程。",
        "基于上述背景，本文围绕牧场牲畜计数与行为监测需求，设计了一种融合YOLOv8与KNN的检测方法，并进一步实现Web监控平台。该研究既具有一定的工程应用价值，也可为智慧牧场中视觉感知、行为分析和异常预警提供参考。"
    ]
    for text in paragraphs:
        add_paragraph(doc, text)
        md.append(text + "\n\n")
    add_picture_placeholder(doc, "图1-1", "智慧牧场监控应用场景图")
    md.append("【图1-1 此处插入智慧牧场监控应用场景图】\n\n")

    add_heading(doc, "1.2 牧场牲畜计数与行为监测研究内容与现状", 2)
    md.append(md_heading("1.2 牧场牲畜计数与行为监测研究内容与现状", 2))
    add_heading(doc, "1.2.1 牧场牲畜计数与行为监测研究发展历程", 3)
    md.append(md_heading("1.2.1 牧场牲畜计数与行为监测研究发展历程", 3))
    paragraphs = [
        "牧场牲畜监测技术经历了从人工观察、传感器监测到视觉智能分析的发展过程。早期牧场管理主要依靠饲养员巡查和人工记录，虽然实现成本较低，但难以保证数据连续性和客观性。随后，一些研究和应用开始使用电子耳标、项圈传感器、计步器和定位设备采集牲畜活动信息，这类方法能够记录部分运动和位置数据，但设备需要佩戴在牲畜身上，存在成本、维护和舒适性问题。",
        "随着摄像头和图像处理技术的发展，基于视频的非接触式监测逐渐成为重要方向。传统图像处理方法通常通过背景建模、边缘检测、颜色分割或人工特征进行目标提取，但牧场环境存在光照变化、遮挡、粪污、栏杆和多目标聚集等干扰，传统方法鲁棒性不足。深度学习目标检测算法出现后，牲畜检测与计数的准确性和适应性显著提高，视觉监控逐渐具备在真实牧场场景中落地的条件。",
        "在行为识别方面，早期方法多依赖人工特征和传统分类器，例如基于纹理、轮廓、姿态比例或运动轨迹特征判断站立、卧躺、行走等行为。近年来，卷积神经网络、循环神经网络和多目标跟踪方法被用于行为识别任务，能够更好地利用空间特征和时序信息。不过，对于本科毕业设计和中小规模场景而言，采用深度检测模型完成目标定位，再结合传统分类器完成基础行为判断，是一种实现难度适中且具有可解释性的方案。"
    ]
    for text in paragraphs:
        add_paragraph(doc, text)
        md.append(text + "\n\n")

    add_heading(doc, "1.2.2 国内外研究现状", 3)
    md.append(md_heading("1.2.2 国内外研究现状", 3))
    paragraphs = [
        "目标检测算法是牲畜计数与行为监测的基础。两阶段检测算法以R-CNN、Fast R-CNN和Faster R-CNN为代表，通常具有较高的检测精度，但模型结构复杂、推理速度较慢。单阶段检测算法以SSD和YOLO系列为代表，将目标类别与位置回归合并在一次前向传播中完成，具备更好的实时性。YOLOv8在网络结构、训练策略和工程封装方面较为成熟，因此适合用于实时监控和部署验证。",
        "在牛只检测和计数研究中，国内外学者普遍关注复杂养殖环境下的遮挡、目标密集、光照变化和视角差异等问题。部分研究通过改进注意力机制、特征融合结构或多目标跟踪算法提升检测稳定性，也有研究将检测结果与跟踪ID结合以减少重复计数。总体来看，单纯依赖检测框可以完成静态图像计数，但在视频场景中仍需要结合目标跟踪和连续帧信息，才能进一步提高统计结果的稳定性。",
        "在牛只行为识别研究中，常见识别目标包括卧躺、站立、行走、采食和饮水等。对于卧躺、站立、行走等基础姿态行为，目标轮廓、长宽比、纹理方向和运动幅度是较有效的判断依据；对于采食、饮水等语义行为，单帧姿态往往不足以准确判断，还需要结合区域位置和持续时间等上下文信息。因此，本文在KNN姿态分类基础上增加区域辅助规则和视频运动修正，以提高行为判断的合理性。",
        "现有智能牧场平台通常包括实时监控、设备管理、异常告警和历史数据分析等功能，但许多研究停留在算法验证阶段，缺少面向实际使用流程的系统实现。本文在算法设计之外进一步搭建Web平台，使识别结果能够保存为行为事件，并通过告警和历史分析功能服务于牧场管理。"
    ]
    for text in paragraphs:
        add_paragraph(doc, text)
        md.append(text + "\n\n")

    add_heading(doc, "1.3 论文结构安排", 2)
    md.append(md_heading("1.3 论文结构安排", 2))
    structure = (
        "本文共分为六章。第1章介绍课题研究背景、研究意义和国内外研究现状。第2章介绍目标检测、卷积神经网络、YOLOv8和KNN行为分类等相关理论。"
        "第3章阐述改进的YOLOv8与KNN融合检测方法，包括迁移学习、行为数据集构建、视频动量修正和区域辅助行为修正。"
        "第4章给出实验环境、数据集、评价指标和实验结果，并对检测模型、行为分类模型以及融合视频检测效果进行分析。"
        "第5章介绍牧场牲畜行为检测系统的设计与实现，包括系统需求、总体架构、功能模块和运行测试。第6章总结全文工作，并对后续研究方向进行展望。"
    )
    add_paragraph(doc, structure)
    md.append(structure + "\n\n")

    # Chapter 2
    add_heading(doc, "第 2 章 目标检测与行为识别相关理论基础", 1)
    md.append(md_heading("第 2 章 目标检测与行为识别相关理论基础", 1))
    add_heading(doc, "2.1 目标检测技术发展概述", 2)
    md.append(md_heading("2.1 目标检测技术发展概述", 2))
    add_heading(doc, "2.1.1 传统目标检测方法", 3)
    md.append(md_heading("2.1.1 传统目标检测方法", 3))
    paragraphs = [
        "传统目标检测方法通常由候选区域提取、人工特征设计和分类器判别三部分组成。常见特征包括Haar特征、HOG特征、SIFT特征等，分类器可采用SVM、Adaboost或KNN等方法。这类方法在背景简单、目标形态稳定的场景中具有一定效果，但对于牧场中光照变化、目标遮挡和姿态差异较大的情况，人工特征往往难以完整描述目标外观。",
        "传统方法的优势在于模型结构简单、计算开销较低、解释性较强；不足在于特征表达能力有限，依赖人工经验较多。当目标尺度变化明显或背景复杂时，检测结果容易出现漏检和误检。因此，在本文中，传统特征方法主要用于行为分类阶段的HOG特征提取，而目标检测部分采用深度学习方法完成。"
    ]
    for text in paragraphs:
        add_paragraph(doc, text)
        md.append(text + "\n\n")

    add_heading(doc, "2.1.2 基于深度学习的目标检测方法", 3)
    md.append(md_heading("2.1.2 基于深度学习的目标检测方法", 3))
    paragraphs = [
        "深度学习目标检测方法通过卷积神经网络自动学习图像特征，减少了人工设计特征的依赖。按照检测流程不同，深度学习目标检测方法通常可分为两阶段算法和单阶段算法。两阶段算法先生成候选区域，再对候选区域进行分类和边界框回归，代表方法包括Faster R-CNN等。该类方法检测精度较高，但推理流程较长。",
        "单阶段算法直接在特征图上预测目标类别和位置，省去了候选区域生成步骤，因此检测速度更快。SSD和YOLO系列算法是单阶段检测的典型代表。在牧场视频监控任务中，系统不仅需要较高准确率，还需要满足实时处理和部署方便的要求，因此单阶段检测算法更加适合本文任务。"
    ]
    for text in paragraphs:
        add_paragraph(doc, text)
        md.append(text + "\n\n")

    add_heading(doc, "2.1.3 单阶段目标检测算法", 3)
    md.append(md_heading("2.1.3 单阶段目标检测算法", 3))
    paragraphs = [
        "单阶段目标检测算法将目标定位和类别判断视为统一回归问题，在一次网络前向传播中输出检测结果。YOLO系列算法将图像划分为网格或基于特征图预测目标框，通过置信度、类别概率和边界框回归得到最终检测结果。其优势是速度快、模型部署成熟，适合视频流场景。",
        "在本文中，YOLOv8被用于牛只目标检测。系统利用检测框数量完成图像中的牛只计数，并将检测框裁剪结果作为后续KNN行为识别的输入，从而形成检测与分类的串联流程。"
    ]
    for text in paragraphs:
        add_paragraph(doc, text)
        md.append(text + "\n\n")
    add_picture_placeholder(doc, "图2-1", "目标检测算法流程对比图")
    md.append("【图2-1 此处插入目标检测算法流程对比图】\n\n")

    add_heading(doc, "2.2 卷积神经网络基础", 2)
    md.append(md_heading("2.2 卷积神经网络基础", 2))
    for title, body in [
        ("2.2.1 卷积层", "卷积层是卷积神经网络的核心结构，通过卷积核在输入特征图上滑动并进行加权求和，提取局部纹理、边缘和形状等信息。不同卷积核可以学习不同类型的特征，深层网络则能够逐步形成更抽象的语义表示。"),
        ("2.2.2 池化层", "池化层用于降低特征图尺寸，减少计算量并提升特征的平移不变性。常见池化方式包括最大池化和平均池化。目标检测网络中也常通过下采样卷积实现类似效果，使网络能够在不同尺度上提取目标特征。"),
        ("2.2.3 激活函数", "激活函数为神经网络引入非线性表达能力。若没有激活函数，多层网络本质上仍可等价为线性变换。常见激活函数包括ReLU、SiLU等，YOLO系列网络中常采用非线性激活提高模型拟合能力。"),
        ("2.2.4 全连接层", "全连接层用于将特征向量映射到输出空间，在分类网络中较为常见。现代目标检测网络通常更多使用卷积结构完成空间预测，以保留目标位置信息并提高推理效率。"),
    ]:
        add_heading(doc, title, 3)
        md.append(md_heading(title, 3))
        add_paragraph(doc, body)
        md.append(body + "\n\n")
    add_picture_placeholder(doc, "图2-2", "卷积神经网络基本结构示意图")
    md.append("【图2-2 此处插入卷积神经网络基本结构示意图】\n\n")

    add_heading(doc, "2.3 YOLOv8 算法原理", 2)
    md.append(md_heading("2.3 YOLOv8 算法原理", 2))
    for title, body in [
        ("2.3.1 YOLOv8 网络结构", "YOLOv8是一种单阶段目标检测算法，整体结构通常由Backbone、Neck和Head三部分组成。Backbone用于提取多层次图像特征，Neck用于进行多尺度特征融合，Head用于输出目标类别、边界框和置信度。"),
        ("2.3.2 Backbone 主干网络", "Backbone主干网络负责从输入图像中提取基础视觉特征。随着网络层数加深，特征图空间尺寸逐渐减小，语义表达能力逐渐增强。对于牛只检测任务，主干网络需要提取牛体轮廓、颜色纹理和姿态等关键信息。"),
        ("2.3.3 Neck 特征融合网络", "Neck特征融合网络负责整合不同尺度的特征信息。牧场视频中牛只可能存在远近差异和尺度变化，多尺度特征融合有助于提高小目标和遮挡目标的检测效果。"),
        ("2.3.4 Head 检测头", "Head检测头根据融合后的特征图输出最终检测结果。检测结果通常包括目标框坐标、类别概率和置信度。本文仅关注牛只类别，因此检测头输出的目标框数量可直接用于计数，并为后续行为分类提供裁剪区域。"),
    ]:
        add_heading(doc, title, 3)
        md.append(md_heading(title, 3))
        add_paragraph(doc, body)
        md.append(body + "\n\n")
    add_picture_placeholder(doc, "图2-3", "YOLOv8网络结构示意图")
    md.append("【图2-3 此处插入YOLOv8网络结构示意图】\n\n")

    add_heading(doc, "2.4 KNN 行为分类算法原理", 2)
    md.append(md_heading("2.4 KNN 行为分类算法原理", 2))
    for title, body in [
        ("2.4.1 KNN 算法基本思想", "KNN算法是一种基于实例的分类方法。对于待分类样本，算法计算其与训练集中各样本的距离，选择距离最近的K个样本，并根据这些邻近样本的类别投票得到最终类别。KNN无需复杂训练过程，适合在特征维度可控、类别数量较少的场景下使用。"),
        ("2.4.2 距离度量方法", "KNN分类结果依赖样本之间的距离度量。常用距离包括欧氏距离、曼哈顿距离和余弦距离。本文在HOG特征归一化后使用欧氏距离衡量样本相似度，并采用距离倒数作为投票权重，使距离更近的样本对分类结果产生更大影响。"),
        ("2.4.3 K 值选择对分类结果的影响", "K值过小容易受到噪声样本影响，K值过大则可能削弱局部类别差异。本文结合验证结果和实现复杂度，采用k=9作为行为分类模型参数，在卧躺、站立和行走三类行为上取得了较稳定的分类效果。"),
    ]:
        add_heading(doc, title, 3)
        md.append(md_heading(title, 3))
        add_paragraph(doc, body)
        md.append(body + "\n\n")
    add_picture_placeholder(doc, "图2-4", "KNN分类原理示意图")
    md.append("【图2-4 此处插入KNN分类原理示意图】\n\n")
    add_heading(doc, "2.5 本章小结", 2)
    md.append(md_heading("2.5 本章小结", 2))
    text = "本章介绍了目标检测、卷积神经网络、YOLOv8算法和KNN行为分类算法的基本原理，为后续融合检测方法设计奠定理论基础。YOLOv8适合承担牛只检测与计数任务，KNN则可结合HOG姿态特征完成基础行为分类。"
    add_paragraph(doc, text)
    md.append(text + "\n\n")

    # Chapter 3
    add_heading(doc, "第 3 章 改进的 YOLOv8 与 KNN 融合检测方法", 1)
    md.append(md_heading("第 3 章 改进的 YOLOv8 与 KNN 融合检测方法", 1))
    add_heading(doc, "3.1 整体算法流程", 2)
    md.append(md_heading("3.1 整体算法流程", 2))
    paragraphs = [
        "本文提出的融合检测方法以视频或图像作为输入，首先调用训练后的YOLOv8模型检测画面中的牛只目标，并得到每个目标的边界框、置信度和目标数量。随后，根据检测框从原始图像中裁剪单牛区域，对裁剪图像进行尺寸归一化和HOG特征提取，再输入KNN分类器得到卧躺、站立或行走行为类别。",
        "在视频场景中，仅依赖单帧图像进行行为识别容易受到姿态瞬间变化、遮挡和图像模糊影响。为提高视频检测稳定性，本文使用YOLOv8的跟踪结果获取同一目标在连续帧中的中心点位置，并根据中心点位移计算运动分数。当运动分数较高时，将行为修正为行走；当运动分数较低且KNN输出行走时，将其修正为站立，从而降低静止目标被误判为行走的概率。",
        "此外，牧场中不同区域具有明确语义，例如采食区、饮水区和休息区。本文在系统中支持人工标注区域，并在推理结果中结合目标底部中心点与区域多边形的位置关系，对行为结果进行二次修正，使检测结果更符合实际场景含义。"
    ]
    for text in paragraphs:
        add_paragraph(doc, text)
        md.append(text + "\n\n")
    add_picture_placeholder(doc, "图3-1", "YOLOv8与KNN融合检测整体流程图")
    md.append("【图3-1 此处插入YOLOv8与KNN融合检测整体流程图】\n\n")

    add_heading(doc, "3.2 基于 YOLOv8 的牲畜目标检测改进方法", 2)
    md.append(md_heading("3.2 基于 YOLOv8 的牲畜目标检测改进方法", 2))
    add_heading(doc, "3.2.1 迁移学习训练策略", 3)
    md.append(md_heading("3.2.1 迁移学习训练策略", 3))
    paragraphs = [
        "直接使用通用目标检测模型处理牧场牛只视频时，模型对牧场场景中的牛只尺度、遮挡和背景干扰适应不足。为提升检测效果，本文采用迁移学习策略：先利用通用动物相关数据训练基础模型，再在奶牛数据集上继续训练，使模型从通用目标特征逐步迁移到牧场牛只检测任务。",
        "实验中采用YOLOv8n作为基线模型，输入尺寸设置为832，使用奶牛数据集进行训练。与从零开始直接训练相比，迁移学习能够利用已有权重中的通用视觉特征，减少训练难度并提升收敛效果。后续实验表明，迁移学习方案在mAP50、Precision和Recall等指标上均优于直接训练方案。"
    ]
    for text in paragraphs:
        add_paragraph(doc, text)
        md.append(text + "\n\n")
    add_picture_placeholder(doc, "图3-2", "YOLOv8迁移学习训练流程图")
    md.append("【图3-2 此处插入YOLOv8迁移学习训练流程图】\n\n")

    add_heading(doc, "3.2.2 检测结果计数方法", 3)
    md.append(md_heading("3.2.2 检测结果计数方法", 3))
    text = "在图像检测场景中，模型输出的有效检测框数量即为当前画面中识别到的牛只数量。系统根据置信度阈值过滤低可信检测框，并通过非极大值抑制去除重复框。在视频场景中，为减少重复计数和检测框抖动带来的影响，系统结合跟踪ID维护目标连续性，使同一牛只在连续帧中保持相对稳定的身份标识。"
    add_paragraph(doc, text)
    md.append(text + "\n\n")
    add_picture_placeholder(doc, "图3-3", "基于检测框的牛只计数示意图")
    md.append("【图3-3 此处插入基于检测框的牛只计数示意图】\n\n")

    add_heading(doc, "3.3 基于 KNN 的牲畜行为识别改进方法", 2)
    md.append(md_heading("3.3 基于 KNN 的牲畜行为识别改进方法", 2))
    add_heading(doc, "3.3.1 行为数据集构建与样本裁剪", 3)
    md.append(md_heading("3.3.1 行为数据集构建与样本裁剪", 3))
    text = "为避免整幅图像中的背景区域干扰行为识别，本文先利用标注框或检测框裁剪出单牛图像，再按照lying、standing和walking三类行为构建数据集。裁剪后的样本能够突出牛只姿态特征，使KNN分类器更加关注目标本身的轮廓、方向和纹理变化。最终数据集划分为训练集、验证集和测试集，其中训练集5363张、验证集666张、测试集648张。"
    add_paragraph(doc, text)
    md.append(text + "\n\n")
    add_picture_placeholder(doc, "图3-4", "单牛行为样本裁剪示意图")
    md.append("【图3-4 此处插入单牛行为样本裁剪示意图】\n\n")

    add_heading(doc, "3.3.2 HOG 特征与 KNN 分类模型", 3)
    md.append(md_heading("3.3.2 HOG 特征与 KNN 分类模型", 3))
    paragraphs = [
        "裁剪图像首先按比例缩放并填充至64×64尺寸，再转换为灰度图并进行高斯滤波。随后使用HOG描述子提取方向梯度直方图特征，用于描述牛只轮廓和姿态方向。为了增强对卧躺与站立等姿态差异的表达，本文还将裁剪区域的长宽比作为补充特征，与HOG特征拼接形成最终特征向量。",
        "KNN分类器训练时保存训练样本特征、标签、特征均值和标准差。推理阶段先对输入特征进行同样的归一化处理，再计算其与训练样本之间的欧氏距离。本文采用k=9，并根据距离倒数进行加权投票，得到行为类别及置信度。"
    ]
    for text in paragraphs:
        add_paragraph(doc, text)
        md.append(text + "\n\n")
    add_picture_placeholder(doc, "图3-5", "HOG特征与KNN行为分类流程图")
    md.append("【图3-5 此处插入HOG特征与KNN行为分类流程图】\n\n")

    add_heading(doc, "3.4 面向视频场景的行为分类优化方法", 2)
    md.append(md_heading("3.4 面向视频场景的行为分类优化方法", 2))
    add_heading(doc, "3.4.1 基于连续帧的运动状态判断", 3)
    md.append(md_heading("3.4.1 基于连续帧的运动状态判断", 3))
    text = "在视频检测中，系统为每个跟踪目标维护最近若干帧的中心点队列，默认窗口长度为6帧。设相邻两帧中心点为(xi, yi)和(xi+1, yi+1)，可计算其欧氏位移，并将多帧平均位移除以目标框宽高中的较大值，得到归一化运动分数。该分数能够反映目标在短时间内的相对运动程度，减少由于摄像头分辨率或目标尺度不同造成的影响。"
    add_paragraph(doc, text)
    md.append(text + "\n\n")

    add_heading(doc, "3.4.2 引入动量信息的行为结果修正", 3)
    md.append(md_heading("3.4.2 引入动量信息的行为结果修正", 3))
    text = "当运动分数高于高阈值时，说明目标在连续帧中发生明显位移，系统将最终行为修正为walking；当运动分数低于低阈值时，说明目标基本处于静止状态，若KNN输出walking，则修正为standing，若KNN输出lying则保持lying。对于运动分数处于中间范围的情况，系统保留KNN分类结果。通过该规则，视频中静止站立牛只被误判为行走的情况可以得到一定缓解。"
    add_paragraph(doc, text)
    md.append(text + "\n\n")
    add_picture_placeholder(doc, "图3-6", "连续帧运动分数与动量修正示意图")
    md.append("【图3-6 此处插入连续帧运动分数与动量修正示意图】\n\n")

    add_heading(doc, "3.5 YOLOv8 与 KNN 融合检测流程", 2)
    md.append(md_heading("3.5 YOLOv8 与 KNN 融合检测流程", 2))
    add_heading(doc, "3.5.1 检测框裁剪与行为分类融合", 3)
    md.append(md_heading("3.5.1 检测框裁剪与行为分类融合", 3))
    text = "融合流程中，YOLOv8负责输出牛只检测框，KNN负责对检测框对应的单牛图像进行行为分类。系统将最终结果以检测框、行为标签、检测置信度、分类置信度和跟踪ID的形式叠加到图像或视频帧上。该方式既保留目标检测的位置优势，又补充了行为识别能力。"
    add_paragraph(doc, text)
    md.append(text + "\n\n")

    add_heading(doc, "3.5.2 基于区域信息的行为结果修正", 3)
    md.append(md_heading("3.5.2 基于区域信息的行为结果修正", 3))
    text = "在实际牧场中，采食、饮水和休息等行为往往与区域位置相关。系统支持在视频画面中标注feeding、water和rest等区域，并以牛只检测框底部中心点作为位置锚点。当目标锚点落入相应区域，且停留时间达到设定阈值时，系统可将基础行为进一步修正为feeding、drinking或resting等语义行为。区域规则作为第二阶段修正，不替代KNN姿态分类，而是增强结果的场景解释能力。"
    add_paragraph(doc, text)
    md.append(text + "\n\n")
    add_picture_placeholder(doc, "图3-7", "区域辅助行为修正示意图")
    md.append("【图3-7 此处插入区域辅助行为修正示意图】\n\n")
    add_heading(doc, "3.6 本章小结", 2)
    md.append(md_heading("3.6 本章小结", 2))
    text = "本章提出了改进的YOLOv8与KNN融合检测方法。该方法通过迁移学习提升牛只检测效果，通过单牛裁剪与HOG+KNN实现基础行为识别，并针对视频场景加入动量修正和区域辅助规则，使检测结果更适合实际牧场监控应用。"
    add_paragraph(doc, text)
    md.append(text + "\n\n")

    # Chapter 4
    add_heading(doc, "第 4 章 实验结果与分析", 1)
    md.append(md_heading("第 4 章 实验结果与分析", 1))
    add_heading(doc, "4.1 前言", 2)
    md.append(md_heading("4.1 前言", 2))
    text = "为验证本文方法的有效性，本章分别对YOLOv8牛只目标检测模型、KNN行为分类模型以及融合视频检测流程进行实验分析。实验重点关注检测精度、行为分类准确率以及视频场景中动量修正对分类稳定性的影响。"
    add_paragraph(doc, text)
    md.append(text + "\n\n")
    add_heading(doc, "4.2 实验环境", 2)
    md.append(md_heading("4.2 实验环境", 2))
    text = "实验在Windows环境下完成，主要使用Python、OpenCV、NumPy和Ultralytics YOLO框架进行模型训练与推理。系统平台部分使用Vue、FastAPI、PostgreSQL、Redis和Docker Compose进行搭建。YOLOv8检测模型采用yolov8n结构，KNN行为分类模型采用自定义NumPy实现，输入图像尺寸为64×64，邻居数k设置为9。"
    add_paragraph(doc, text)
    md.append(text + "\n\n")

    add_heading(doc, "4.3 数据集构建", 2)
    md.append(md_heading("4.3 数据集构建", 2))
    add_heading(doc, "4.3.1 数据集介绍", 3)
    md.append(md_heading("4.3.1 数据集介绍", 3))
    text = "本文实验数据主要包括牛只目标检测数据集和牛只行为分类数据集。检测数据集用于训练YOLOv8模型，使其能够定位图像中的牛只目标。行为分类数据集由单牛裁剪图像组成，按照lying、standing和walking三类行为进行组织，用于训练和测试KNN分类器。"
    add_paragraph(doc, text)
    md.append(text + "\n\n")
    add_picture_placeholder(doc, "图4-1", "检测数据集与行为数据集样本示例")
    md.append("【图4-1 此处插入检测数据集与行为数据集样本示例】\n\n")
    add_heading(doc, "4.3.2 数据集划分", 3)
    md.append(md_heading("4.3.2 数据集划分", 3))
    text = "行为分类数据集划分为训练集、验证集和测试集。训练集共5363张，其中lying 1909张、standing 2540张、walking 914张；验证集共666张，其中lying 232张、standing 324张、walking 110张；测试集共648张，其中lying 230张、standing 298张、walking 120张。"
    add_paragraph(doc, text)
    md.append(text + "\n\n")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ["数据集", "lying", "standing", "walking"]):
        set_cell_text(cell, text)
    for row in [
        ("训练集", "1909", "2540", "914"),
        ("验证集", "232", "324", "110"),
        ("测试集", "230", "298", "120"),
    ]:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            set_cell_text(cell, text)
    md.append("| 数据集 | lying | standing | walking |\n|---|---:|---:|---:|\n| 训练集 | 1909 | 2540 | 914 |\n| 验证集 | 232 | 324 | 110 |\n| 测试集 | 230 | 298 | 120 |\n\n")
    add_picture_placeholder(doc, "图4-2", "行为数据集划分统计图")
    md.append("【图4-2 此处插入行为数据集划分统计图】\n\n")

    add_heading(doc, "4.4 评价指标", 2)
    md.append(md_heading("4.4 评价指标", 2))
    paragraphs = [
        "目标检测实验主要采用Precision、Recall、mAP50和mAP50-95作为评价指标。Precision表示检测为正样本的目标中真正为牛只的比例，Recall表示真实牛只目标被成功检测出的比例。mAP50表示IoU阈值为0.5时的平均精度，mAP50-95则是在多个IoU阈值下计算的平均精度，更能反映模型综合定位能力。",
        "行为分类实验采用整体准确率、各类别准确率和混淆矩阵进行评价。整体准确率反映模型在全部测试样本上的平均分类正确率，类别准确率用于观察不同姿态行为的识别差异，混淆矩阵则能够直观展示不同类别之间的误判情况。"
    ]
    for text in paragraphs:
        add_paragraph(doc, text)
        md.append(text + "\n\n")

    add_heading(doc, "4.5 实验结果与分析", 2)
    md.append(md_heading("4.5 实验结果与分析", 2))
    add_heading(doc, "4.5.1 YOLOv8 检测实验结果", 3)
    md.append(md_heading("4.5.1 YOLOv8 检测实验结果", 3))
    text = "本文对不同训练策略进行了对比。通用基础模型在验证集上的mAP50为0.7134，直接在奶牛数据集训练后的mAP50提升至0.8560，而采用迁移学习继续训练后的mAP50达到0.9018，mAP50-95达到0.7094。结果说明，迁移学习能够有效提升模型对奶牛场景的适应能力。"
    add_paragraph(doc, text)
    md.append(text + "\n\n")
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ["训练策略", "Epoch", "Precision", "Recall", "mAP50", "mAP50-95"]):
        set_cell_text(cell, text)
    for row in [
        ("通用基础模型", "105", "0.6894", "0.6808", "0.7134", "0.6035"),
        ("奶牛数据集直接训练", "150", "0.8547", "0.7786", "0.8560", "0.6289"),
        ("迁移学习训练", "120", "0.9095", "0.8292", "0.9018", "0.7094"),
    ]:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            set_cell_text(cell, text)
    md.append("| 训练策略 | Epoch | Precision | Recall | mAP50 | mAP50-95 |\n|---|---:|---:|---:|---:|---:|\n| 通用基础模型 | 105 | 0.6894 | 0.6808 | 0.7134 | 0.6035 |\n| 奶牛数据集直接训练 | 150 | 0.8547 | 0.7786 | 0.8560 | 0.6289 |\n| 迁移学习训练 | 120 | 0.9095 | 0.8292 | 0.9018 | 0.7094 |\n\n")
    add_picture_placeholder(doc, "图4-3", "YOLOv8训练结果曲线图")
    md.append("【图4-3 此处插入YOLOv8训练结果曲线图】\n\n")
    add_picture_placeholder(doc, "图4-4", "YOLOv8检测模型混淆矩阵或检测效果图")
    md.append("【图4-4 此处插入YOLOv8检测模型混淆矩阵或检测效果图】\n\n")

    add_heading(doc, "4.5.2 KNN 行为识别实验结果", 3)
    md.append(md_heading("4.5.2 KNN 行为识别实验结果", 3))
    text = "KNN行为分类模型在验证集上的整体准确率为83.03%，在测试集上的整体准确率为87.96%。从测试集结果看，lying类别准确率为92.17%，standing类别准确率为83.22%，walking类别准确率为91.67%。standing类别相对较低，主要原因是站立与行走在单帧图像中姿态相似，且部分过渡帧难以仅凭静态特征准确区分。"
    add_paragraph(doc, text)
    md.append(text + "\n\n")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ["数据集", "整体准确率", "lying", "standing / walking"]):
        set_cell_text(cell, text)
    for row in [
        ("验证集", "0.8303", "0.9052", "standing 0.7623；walking 0.8727"),
        ("测试集", "0.8796", "0.9217", "standing 0.8322；walking 0.9167"),
    ]:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            set_cell_text(cell, text)
    md.append("| 数据集 | 整体准确率 | lying | standing / walking |\n|---|---:|---:|---|\n| 验证集 | 0.8303 | 0.9052 | standing 0.7623；walking 0.8727 |\n| 测试集 | 0.8796 | 0.9217 | standing 0.8322；walking 0.9167 |\n\n")
    add_picture_placeholder(doc, "图4-5", "KNN行为分类混淆矩阵")
    md.append("【图4-5 此处插入KNN行为分类混淆矩阵】\n\n")

    add_heading(doc, "4.5.3 融合模型视频检测与动量修正结果分析", 3)
    md.append(md_heading("4.5.3 融合模型视频检测与动量修正结果分析", 3))
    paragraphs = [
        "融合模型能够在视频帧中同时显示牛只检测框、跟踪ID、行为类别、KNN置信度和YOLO检测置信度。对于单帧图像，系统主要依赖KNN输出行为类别；对于视频流，系统进一步利用连续帧位移判断目标运动状态，使行为结果更加稳定。",
        "在未加入动量修正时，部分静止站立牛只可能由于姿态倾斜或局部模糊被识别为walking。加入动量修正后，若目标在连续帧中的归一化运动分数较低，则系统会将walking修正为standing；若运动分数较高，则将目标判断为walking。该策略利用视频时序信息弥补了单帧分类的不足，使站立与行走两类行为的区分更加符合实际观察结果。",
        "需要注意的是，动量修正规则依赖稳定的目标跟踪ID。当发生严重遮挡或跟踪ID切换时，连续帧中心点历史可能被打断，系统会暂时回退到KNN单帧分类结果。因此，后续仍可进一步引入更稳定的多目标跟踪算法或时序行为识别模型。"
    ]
    for text in paragraphs:
        add_paragraph(doc, text)
        md.append(text + "\n\n")
    add_picture_placeholder(doc, "图4-6", "融合模型视频检测结果截图")
    md.append("【图4-6 此处插入融合模型视频检测结果截图】\n\n")
    add_picture_placeholder(doc, "图4-7", "动量修正前后分类结果对比图")
    md.append("【图4-7 此处插入动量修正前后分类结果对比图】\n\n")

    add_heading(doc, "4.6 本章小结", 2)
    md.append(md_heading("4.6 本章小结", 2))
    text = "本章对本文方法进行了实验验证。结果表明，迁移学习训练能够有效提升YOLOv8在奶牛检测任务中的表现；HOG+KNN能够完成基础行为分类；视频动量修正能够改善单帧分类在站立与行走行为上的部分误判，为系统应用提供了实验支撑。"
    add_paragraph(doc, text)
    md.append(text + "\n\n")

    # Chapter 5
    add_heading(doc, "第 5 章 牧场牲畜行为检测系统设计与实现", 1)
    md.append(md_heading("第 5 章 牧场牲畜行为检测系统设计与实现", 1))
    add_heading(doc, "5.1 前言", 2)
    md.append(md_heading("5.1 前言", 2))
    text = "在完成算法设计与实验验证后，本文进一步设计并实现牧场牲畜行为检测系统，使检测结果能够通过平台进行展示、存储和分析。本章主要介绍系统需求、总体设计、功能模块实现以及运行测试情况。"
    add_paragraph(doc, text)
    md.append(text + "\n\n")
    add_heading(doc, "5.2 系统需求分析", 2)
    md.append(md_heading("5.2 系统需求分析", 2))
    paragraphs = [
        "系统面向牧场管理人员和监控值班人员，主要目标是降低人工巡检压力，提升牛只状态观察和异常处理效率。系统需要支持用户登录、设备管理、实时监控、区域标注、行为事件记录、告警规则配置和历史数据分析等功能。",
        "从功能需求看，系统应能够接入多个监控设备，展示设备状态和视频画面；能够调用推理服务生成行为事件，并将事件写入数据库；能够根据用户配置的规则生成告警；能够按时间、设备和行为类型查询历史记录。从非功能需求看，系统应具备较好的易用性、可扩展性和部署便利性。"
    ]
    for text in paragraphs:
        add_paragraph(doc, text)
        md.append(text + "\n\n")

    add_heading(doc, "5.3 系统总体设计", 2)
    md.append(md_heading("5.3 系统总体设计", 2))
    add_heading(doc, "5.3.1 系统总体架构", 3)
    md.append(md_heading("5.3.1 系统总体架构", 3))
    text = "系统采用前后端分离架构。前端基于Vue和TypeScript实现页面展示与交互；后端基于FastAPI提供业务接口；推理服务独立部署，负责处理图像、视频或视频流输入并返回行为事件；PostgreSQL用于存储用户、设备、区域、事件和告警数据；Redis用于缓存和辅助状态管理；MediaMTX用于提供RTSP和HLS流媒体能力；Docker Compose用于统一编排各服务。"
    add_paragraph(doc, text)
    md.append(text + "\n\n")
    add_picture_placeholder(doc, "图5-1", "系统总体架构图")
    md.append("【图5-1 此处插入系统总体架构图】\n\n")

    add_heading(doc, "5.3.2 系统功能模块设计", 3)
    md.append(md_heading("5.3.2 系统功能模块设计", 3))
    text = "系统功能模块包括用户登录模块、设备管理模块、实时监控模块、区域配置模块、行为事件模块、告警模块和历史分析模块。各模块通过统一API与后端交互，后端根据业务逻辑访问数据库或调用推理服务，从而完成数据展示、事件生成和告警管理。"
    add_paragraph(doc, text)
    md.append(text + "\n\n")
    add_picture_placeholder(doc, "图5-2", "系统功能模块图")
    md.append("【图5-2 此处插入系统功能模块图】\n\n")

    add_heading(doc, "5.4 系统功能实现", 2)
    md.append(md_heading("5.4 系统功能实现", 2))
    for title, body, fig in [
        ("5.4.1 用户登录模块", "用户登录模块用于完成账号密码登录、JWT会话保持和基础权限控制。系统预设管理员和普通用户两类角色，管理员可维护设备和规则，普通用户主要用于查看数据和管理当前农场范围内的区域。", "图5-3 此处插入用户登录界面截图"),
        ("5.4.2 设备管理模块", "设备管理模块用于维护摄像头等监控设备信息，包括设备编号、设备名称、接入地址、启用状态和所属农场等字段。通过设备管理，系统能够统一管理多路视频来源，为后续实时监控和行为事件生成提供基础。", "图5-4 此处插入设备管理界面截图"),
        ("5.4.3 实时监控模块", "实时监控模块用于展示设备视频画面和当前监控状态。系统通过流媒体服务将视频转换为可在浏览器访问的HLS地址，前端页面负责播放视频流并展示设备在线状态。该模块是牧场管理人员进行日常观察的主要入口。", "图5-5 此处插入实时监控界面截图"),
        ("5.4.4 行为事件与告警模块", "行为事件模块用于记录推理服务返回的行为类型、牛只数量、置信度、设备编号和发生时间等信息。告警模块在行为事件入库后根据预设规则进行判断，当检测到异常停留、异常行为或规则命中时生成告警记录，并支持查看和处理告警状态。", "图5-6 此处插入行为事件与告警中心界面截图"),
        ("5.4.5 历史数据分析模块", "历史数据分析模块支持按设备、行为类型和时间范围查询历史行为事件与告警记录，并通过趋势图和占比图展示行为分布情况。该模块能够帮助管理人员回顾牛群行为变化，为后续养殖管理决策提供数据依据。", "图5-7 此处插入历史数据分析界面截图"),
    ]:
        add_heading(doc, title, 3)
        md.append(md_heading(title, 3))
        add_paragraph(doc, body)
        md.append(body + "\n\n")
        caption, note = fig.split(" 此处插入", 1)
        add_picture_placeholder(doc, caption, note)
        md.append(f"【{fig}】\n\n")

    add_heading(doc, "5.5 系统运行测试", 2)
    md.append(md_heading("5.5 系统运行测试", 2))
    paragraphs = [
        "系统通过Docker Compose启动前端、后端、推理服务、数据库、缓存和流媒体服务。启动后，用户可访问前端页面完成登录，并通过API文档检查后端服务状态。推理服务可接收图片、视频、视频流或边缘设备上报结果，返回统一格式的行为事件。",
        "测试内容主要包括登录状态保持、设备增删改查、区域配置、行为事件导入、告警生成、历史分页查询和数据可视化展示。测试结果表明，系统能够完成从设备管理、视频监控、行为事件生成到告警展示的基本业务闭环。"
    ]
    for text in paragraphs:
        add_paragraph(doc, text)
        md.append(text + "\n\n")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ["测试项", "测试内容", "结果"]):
        set_cell_text(cell, text)
    for row in [
        ("登录模块", "账号密码登录与会话保持", "通过"),
        ("设备模块", "设备列表、新增、编辑和状态展示", "通过"),
        ("行为事件", "推理结果写入与列表展示", "通过"),
        ("告警模块", "规则判断与告警状态更新", "通过"),
        ("历史分析", "按条件查询并展示统计图表", "通过"),
    ]:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            set_cell_text(cell, text)
    md.append("| 测试项 | 测试内容 | 结果 |\n|---|---|---|\n| 登录模块 | 账号密码登录与会话保持 | 通过 |\n| 设备模块 | 设备列表、新增、编辑和状态展示 | 通过 |\n| 行为事件 | 推理结果写入与列表展示 | 通过 |\n| 告警模块 | 规则判断与告警状态更新 | 通过 |\n| 历史分析 | 按条件查询并展示统计图表 | 通过 |\n\n")
    add_picture_placeholder(doc, "图5-8", "系统运行测试截图")
    md.append("【图5-8 此处插入系统运行测试截图】\n\n")

    add_heading(doc, "5.6 本章小结", 2)
    md.append(md_heading("5.6 本章小结", 2))
    text = "本章介绍了牧场牲畜行为检测系统的设计与实现。系统采用前后端分离和独立推理服务架构，完成了设备管理、实时监控、行为事件、告警中心和历史分析等功能，实现了算法结果的平台化展示与管理。"
    add_paragraph(doc, text)
    md.append(text + "\n\n")

    # Chapter 6
    add_heading(doc, "第 6 章 总结与展望", 1)
    md.append(md_heading("第 6 章 总结与展望", 1))
    add_heading(doc, "6.1 总结", 2)
    md.append(md_heading("6.1 总结", 2))
    paragraphs = [
        "本文围绕智能牧场中牲畜计数与行为监测需求，设计并实现了一种融合KNN与YOLOv8的检测方法和Web监控系统。算法方面，本文采用YOLOv8完成牛只目标检测与计数，通过迁移学习提升模型在奶牛场景中的检测效果；在行为识别方面，本文构建单牛行为数据集，利用HOG特征和KNN分类器识别卧躺、站立和行走三类基础行为。",
        "针对视频场景中单帧分类稳定性不足的问题，本文引入连续帧运动信息计算目标运动分数，并通过动量修正规则辅助区分站立与行走状态。同时，系统结合区域标注信息对采食、饮水和休息等语义行为进行辅助判断。实验结果表明，YOLOv8迁移学习模型mAP50达到0.9018，KNN行为分类模型测试集准确率达到87.96%，说明本文方法具有一定有效性。",
        "系统实现方面，本文搭建了牧场牲畜行为检测平台，实现用户登录、设备管理、实时监控、行为事件、告警规则和历史分析等功能，使算法结果能够以更直观的方式服务于牧场管理。"
    ]
    for text in paragraphs:
        add_paragraph(doc, text)
        md.append(text + "\n\n")

    add_heading(doc, "6.2 展望", 2)
    md.append(md_heading("6.2 展望", 2))
    paragraphs = [
        "本文工作仍存在一些不足。首先，KNN行为分类主要依赖单牛裁剪图像的静态特征，对遮挡、模糊和过渡姿态的适应能力有限。后续可引入轻量级卷积神经网络或时序行为识别模型，进一步提升复杂场景下的行为识别准确率。",
        "其次，视频动量修正规则依赖目标跟踪稳定性，当目标严重遮挡或跟踪ID频繁切换时，连续帧运动信息会受到影响。后续可结合更加稳定的多目标跟踪算法，或引入轨迹重识别方法提高目标身份保持能力。",
        "最后，当前系统主要实现基础监控、事件记录和告警展示，后续可继续扩展个体识别、健康趋势评估、移动端通知和边缘设备部署能力，使系统更加贴近真实牧场生产应用。"
    ]
    for text in paragraphs:
        add_paragraph(doc, text)
        md.append(text + "\n\n")

    # References and thanks
    add_heading(doc, "参考文献", 1)
    md.append(md_heading("参考文献", 1))
    refs = [
        "[1] Dalal N, Triggs B. Histograms of oriented gradients for human detection[C]//IEEE Conference on Computer Vision and Pattern Recognition. 2005.",
        "[2] Cover T, Hart P. Nearest neighbor pattern classification[J]. IEEE Transactions on Information Theory, 1967, 13(1): 21-27.",
        "[3] Girshick R. Fast R-CNN[C]//IEEE International Conference on Computer Vision. 2015.",
        "[4] Ren S, He K, Girshick R, et al. Faster R-CNN: Towards real-time object detection with region proposal networks[J]. IEEE Transactions on Pattern Analysis and Machine Intelligence, 2017, 39(6): 1137-1149.",
        "[5] Liu W, Anguelov D, Erhan D, et al. SSD: Single shot multibox detector[C]//European Conference on Computer Vision. 2016.",
        "[6] Redmon J, Divvala S, Girshick R, et al. You Only Look Once: Unified, real-time object detection[C]//IEEE Conference on Computer Vision and Pattern Recognition. 2016.",
        "[7] Redmon J, Farhadi A. YOLO9000: Better, faster, stronger[C]//IEEE Conference on Computer Vision and Pattern Recognition. 2017.",
        "[8] Bochkovskiy A, Wang C Y, Liao H Y M. YOLOv4: Optimal speed and accuracy of object detection[EB/OL]. arXiv:2004.10934, 2020.",
        "[9] Jocher G, Chaurasia A, Qiu J. Ultralytics YOLO[EB/OL]. 2023.",
        "[10] Bewley A, Ge Z, Ott L, et al. Simple online and realtime tracking[C]//IEEE International Conference on Image Processing. 2016.",
        "[11] Wojke N, Bewley A, Paulus D. Simple online and realtime tracking with a deep association metric[C]//IEEE International Conference on Image Processing. 2017.",
        "[12] He K, Zhang X, Ren S, et al. Deep residual learning for image recognition[C]//IEEE Conference on Computer Vision and Pattern Recognition. 2016.",
        "[13] Lin T Y, Dollár P, Girshick R, et al. Feature pyramid networks for object detection[C]//IEEE Conference on Computer Vision and Pattern Recognition. 2017.",
        "[14] Woo S, Park J, Lee J Y, et al. CBAM: Convolutional block attention module[C]//European Conference on Computer Vision. 2018.",
        "[15] Tan M, Pang R, Le Q V. EfficientDet: Scalable and efficient object detection[C]//IEEE Conference on Computer Vision and Pattern Recognition. 2020.",
        "[16] 余涛. 融合KNN与YOLOv8的智能牧场牲畜计数与行为监测系统开题报告[R]. 2026.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(ref)
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(10.5)
        md.append(ref + "\n\n")

    add_heading(doc, "致 谢", 1)
    md.append(md_heading("致 谢", 1))
    thanks = (
        "在本论文完成过程中，指导老师在课题方向、系统设计和论文撰写等方面给予了耐心指导和帮助，使我能够逐步完成从算法实验到系统实现的完整工作。"
        "同时，感谢同学和朋友在资料整理、系统测试和论文修改过程中提供的帮助。通过本次毕业设计，我进一步理解了目标检测、行为识别和Web系统开发的基本流程，也提升了工程实践和问题分析能力。"
        "由于本人水平有限，论文中仍存在不足之处，恳请各位老师批评指正。"
    )
    add_paragraph(doc, thanks)
    md.append(thanks + "\n\n")

    doc.save(DOCX_PATH)
    MD_PATH.write_text("".join(md), encoding="utf-8")


if __name__ == "__main__":
    build_document()
    print(DOCX_PATH)
    print(MD_PATH)
